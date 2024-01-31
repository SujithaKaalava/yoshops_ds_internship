#!/usr/bin/env python
# coding: utf-8
'''Task4: Digital Marketing Automation 
---------------------------
Input Value Type text-  Item name-Data Science Basic Training Program for Everyone(Age=10 to 60)
Input Value Type link  : https://yoshops.com/products/data-science-basic-training-program-for-everyone-age-10-t0-60

Output:

1.Create Banner(Jpg,Png),Tag Line for topic(notepad),Topic description  file(notepad),keyword file(notepad),

2.Create 20sec shorts videos(MP4) on Data Science Basic Training Program for Everyone

3.Write a blog on Data Science Basic Training Program for Everyone.'''



# In[1]:


from PIL import Image, ImageDraw, ImageFont

# Input values
item_name = "Item name-Data Science Basic Training Program for Everyone(Age=10 to 60)"
link = "https://yoshops.com/products/data-science-basic-training-program-for-everyone-age-10-t0-60"

# Load background image for banner
background_image = Image.open(r'C:\Users\Srinu\Desktop\sujitha\pictures\1133470.jpg')

# Set banner height to 1/4 of the background image
banner_height = background_image.height // 4

# Create banner image
banner_image = Image.new('RGBA', (background_image.width, banner_height), color=(0, 0, 0, 126))
draw=ImageDraw.Draw(banner_image )

#font size of banner is to 1/10 of the banner height
font_size=banner_height//8
font = ImageFont.load_default()

#draw item name text
text_width,text_height=draw.textsize(item_name,font=font)
text_x=(banner_image.width-text_width)//2
text_y=(banner_image.height-text_height*2)//2
draw.text((text_x,text_y),item_name,fill=(255,255,255),font=font)

#draw link text
link_font_size=banner_height//8
link_font=ImageFont.load_default()
link_text_width,link_text_height=draw.textsize(link,font=link_font)
link_text_x=(banner_image.width-link_text_width)//2
link_text_y=(banner_image.height-link_text_height)//2+text_height
draw.text((link_text_x,link_text_y),link,(255,255,255),font=link_font)

#set banner position to centre of background image
banner_position=(background_image.width//2-banner_image.width//2,background_image.height//2-banner_height//2)

#overlay banner image on background image
background_image.paste(banner_image,banner_position,banner_image)

#save the banner image
x=background_image.save('project_task4_sol1.jpg','JPEG')


# In[2]:


#creating tag line (notepad)
tagline="Capture Your Thoughts, Unleash Your Creativity with Notepads"
tagline_file_name="project_task4_taglinesol.txt"
with open(tagline_file_name,"w") as file:
    file.write(tagline)
    

#create topic description
topic_description="Notepads are essential tools for capturing thoughts, ideas, and important information. They provide a convenient and tangible way to jot down notes, sketches, to-do lists, and more. With their portable and lightweight design, notepads offer flexibility, allowing you to record your thoughts on the go, whether you're in a meeting, attending a lecture, or simply seeking inspiration"
topic_description_filename="tagline_description.txt"
with open(topic_description_filename,"w") as file:
    file.write(topic_description)

    
    
#keyword file
keywords="Notepad, note-taking, writing, organization, ideas, creativity, thoughts, sketches, jotting down, portable, versatile, productivity, brainstorming, capture, reference, stationery, notebook, paper, inspiration, handy"
keywords_filename="keywords.txt"
with open(keywords_filename,"w") as file:
    file.write(keywords)


# In[ ]:


from moviepy.editor import *

# Load audio file
audio = AudioFileClip(r'E:\MP3 songs\7th Sense\03 - Yellelama [www.AtoZmp3.net]_2_2.mp3')

# Load image file
image = ImageClip(r'C:\Users\Srinu\ds internship\notepad_banner.jpg')

# Set audio for the image clip
image = image.set_audio(audio)

# Set the duration of the image clip to match the audio clip
image = image.set_duration(audio.duration)

# Combine audio and image to create a video clip
video = CompositeVideoClip([image])

# Export the video as an MP4 file
video.write_videofile('task4_video.mp4', fps=30, audio_codec='aac', audio_bitrate='192k')




# In[3]:


# creating blog in word document
import docx


#create a word document
document=docx.Document()

#set the title to the document
document.add_heading("Data Science Basic Training Program for Everyone",0)


# introduction to the document
document.add_paragraph("In today's data-driven world, the demand for skilled data scientists is soaring. From business analytics to healthcare and beyond, the insights derived from data hold the key to solving complex problems and making informed decisions. However, the field of data science can often seem daunting and exclusive to those with specialized backgrounds. That's where the Data Science Basic Training Program for Everyone comes in, offering a comprehensive learning experience designed to empower individuals of all ages and backgrounds to embark on their data science journey")


# add sections for the project
document.add_heading("The Need for Data Science Skills",level=1)
document.add_paragraph("In an era of information overload, the ability to extract meaningful insights from vast amounts of data has become crucial. The Data Science Basic Training Program recognizes this need and aims to bridge the gap by equipping participants with essential data science skills. Regardless of whether you have a background in technology or not, this program opens the doors to the fascinating world of data science.An Inclusive Learning Environment")


#add another section for the document
document.add_heading("Comprehensive Curriculum",level=1)
document.add_paragraph("The program covers a wide range of topics essential for understanding data science. From statistical analysis and machine learning algorithms to data visualization and exploratory data analysis, each aspect is carefully explained and practiced. The curriculum strikes a balance between theoretical concepts and hands-on practical exercises, allowing participants to gain a solid foundation in data science principles")


# adding conclusion  to the document
document.add_heading("conclusion",level=1)
document.add_paragraph("In a world increasingly reliant on data, the Data Science Basic Training Program for Everyone opens the doors to a wealth of opportunities. It enables individuals from diverse backgrounds to develop a strong foundation in data science, unlocking the potential to transform industries and make an impact. With expert guidance, practical projects, and a supportive community, this program paves the way for anyone interested in embarking on a journey into the captivating realm of data science. Embrace the power of data and embark on your data science adventure today!")


#save the document
document.save("task4_blog.docx")

 

